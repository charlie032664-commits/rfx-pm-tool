# -*- coding: utf-8 -*-
"""
LLM Requirements Extractor (Enhanced, overwrite-ready)
- Input : inbound/<case_id>/rfq/*.docx|*.doc|*.xlsx|*.xls|*.pdf|*.md|*.txt
- Output: runs/<case_id>/requirements.json

Features:
- Read DOCX paragraphs + tables (each row -> block)
- Add [ROW_ID=...] hints for stable req_id
- Progress prints per chunk
- Retry with backoff for API errors
- Resume support using runs/<case_id>/requirements.partial.jsonl (jsonl per chunk)
- Fix template req_id like AUTO-<file>-<chunk>-001 (LLM sometimes returns template literally)
- Optional reset partial for clean rerun
- De-dup when loading partial

Usage (run from scripts/ai_rfx):
  python extract_requirements_llm.py --case inbound\\20260129_IBM_RFQ --runs runs --max-chars 1200 --group-size 3 --resume
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple, Set

import httpx
from openai import OpenAI

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from llm_client import get_client, get_model, is_available, parse_json_response
from file_selection import load_excluded   # Phase 7: enforce Step 1.5 exclusion

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None


def analyze_document_schema(client, model: str, rfq_dir: Path) -> dict:
    """多輪分析：先判斷每個檔案格式，再綜合判斷整份 RFQ 格式"""

    def _sample_file(fp: Path) -> str:
        suffix = fp.suffix.lower()
        text = f"=== File: {fp.name} ===\n"
        try:
            if suffix == ".docx" and docx is not None:
                d = docx.Document(str(fp))
                paras = [p.text.strip() for p in d.paragraphs[:60] if p.text.strip()]
                text += "\n".join(paras[:40]) + "\n"
                for ti, table in enumerate(d.tables[:10], 1):
                    text += f"\n--- Table {ti} (first 2 rows) ---\n"
                    for ri, row in enumerate(table.rows[:2], 1):
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        text += f"Row {ri}: {' | '.join(cells)}\n"
                    # 額外輸出第一欄的前3個 data row 值（讓 LLM 識別 req_id 格式）
                    if len(table.rows) > 1:
                        first_col_values = []
                        for row in table.rows[1:4]:
                            val = row.cells[0].text.strip() if row.cells else ""
                            if val:
                                first_col_values.append(val)
                        if first_col_values:
                            text += f"ID column sample values: {' / '.join(first_col_values)}\n"
            elif suffix in (".xlsx", ".xls") and openpyxl is not None:
                wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
                for sn in wb.sheetnames[:4]:
                    ws = wb[sn]
                    rows = list(ws.iter_rows(values_only=True))[:4]
                    text += f"Sheet: {sn}\n"
                    for i, row in enumerate(rows, 1):
                        cells = [str(c) for c in row if c is not None][:8]
                        text += f"  Row {i}: {' | '.join(cells)}\n"
                wb.close()
            elif suffix == ".pdf" and pdfplumber is not None:
                try:
                    with pdfplumber.open(str(fp)) as pdf:
                        text += f"Total pages: {len(pdf.pages)}\n"
                        # 讀前5頁文字
                        for i, page in enumerate(pdf.pages[:5], 1):
                            page_text = (page.extract_text() or "").strip()
                            if page_text:
                                text += f"\n--- Page {i} ---\n"
                                text += page_text[:500] + "\n"
                        # 讀前3個 table 的 header
                        table_count = 0
                        for i, page in enumerate(pdf.pages[:20], 1):
                            tables = page.extract_tables()
                            for t in tables:
                                if t and table_count < 3:
                                    header = [str(c or "").strip()[:20] for c in t[0] if c]
                                    if header:
                                        text += f"\n--- Page {i} Table header: {' | '.join(header)} ---\n"
                                    table_count += 1
                except Exception as e:
                    text += f"(pdf read error: {e})\n"
        except Exception as e:
            text += f"(read error: {e})\n"
        return text

    def _analyze_one_file(fp: Path, sample: str) -> dict:
        prompt = f"""分析以下單一 RFQ 檔案的格式，只輸出 JSON：

{sample[:3000]}

回傳格式：
{{
  "file": "{fp.name}",
  "file_type": "docx/xlsx/pdf",
  "format": "ibm_matrix/simple_list/excel_checklist/plain_text/spec_reference/appendix/unknown",
  "is_main_requirement": true/false,
  "req_id_style": "HOST-N / SYS-N / 4.1.1 / REQ-xxx / none",
  "table_cols": "描述 table 欄位（若有）",
  "confidence": 0.0,
  "notes": ""
}}

注意：
- ibm_matrix: docx 內有 Item | IBM Requirement | IBM Comment | Vendor Comment 四欄
- simple_list: 有編號欄（REQ-xxx, SYS-N）+ 需求內容欄
- excel_checklist: xlsx 內有條列式需求 + 勾選欄
- spec_reference: PDF 產品規格書（SRS/SRD），有 Model | LE1 | LE2 這種多產品規格對照表，is_main_requirement = false
- plain_text: 若 PDF 有明確需求條列（shall/must/required），is_main_requirement = true
- appendix: 問卷、附件、報價單、非需求主文件
- is_main_requirement: 這個檔案是否包含主要需求（非附件/問卷/純規格參考書）

req_id_rule 判斷規則（非常重要，請嚴格遵守）：
- 若 ID column sample values 包含 "HOST N"（如 HOST 1, HOST 2）：
  req_id_rule 必須填：HOST N -> RFQ-HOST-{{N:03d}}, BIOS N -> RFQ-BIOS-{{N:03d}}, BMC N -> RFQ-BMC-{{N:03d}}
- 若 ID column sample values 包含章節編號（如 4.1, 4.2.1, 5.1）：
  req_id_rule 填：SRS Ref.# directly
- 若 ID column sample values 包含 REQ- 開頭的編號：
  req_id_rule 填：REQ-ID directly
- 若版本號（0.1, 1.0）或文字名稱（Altitude, Temperature）：
  這不是 req_id，填：AI auto
- 只有在完全沒有任何編號格式時，才填：AI auto

req_id_rule 只能用 ASCII 字元，不要用 Unicode 箭頭符號（用 -> 而非 ->）。"""

        try:
            resp = client.chat.completions.create(
                model=model, temperature=0,
                messages=[
                    {"role": "system", "content": _SYSTEM_JSON},
                    {"role": "user", "content": prompt},
                ],
            )
            text = (resp.choices[0].message.content or "").strip()
            return parse_json_response(text, model=model)
        except Exception as e:
            print(f"[WARN] analyze file {fp.name}: {str(e).encode('ascii', errors='replace').decode('ascii')[:300]}")
        return {"file": fp.name, "format": "unknown", "confidence": 0}

    # === 第一輪：逐檔分析 ===
    per_file_schemas = []
    for fp in sorted(rfq_dir.glob("*")):
        if fp.suffix.lower() not in (".docx", ".doc", ".xlsx", ".xls", ".pdf"):
            continue
        print(f"[INFO] Analyzing file format: {fp.name}")
        sample = _sample_file(fp)
        file_schema = _analyze_one_file(fp, sample)
        per_file_schemas.append(file_schema)
        safe = str(file_schema).encode("ascii", errors="replace").decode("ascii")
        print(f"[INFO] File schema: {safe}")

    if not per_file_schemas:
        return {}

    # === 第二輪：綜合判斷 ===
    summary = json.dumps(per_file_schemas, ensure_ascii=False, indent=2)
    prompt2 = f"""以下是一份 RFQ 的各個檔案格式分析結果：

{summary}

請綜合判斷整份 RFQ 的格式，只輸出 JSON：
{{
  "customer": "客戶名稱（若看不出則空字串）",
  "rfq_format": "整份 RFQ 的主要格式（以 is_main_requirement=true 的檔案為準）",
  "req_id_style": "主文件的 req_id 格式",
  "req_id_rule": "req_id 轉換規則，使用 ASCII 箭頭 ->，例如：HOST N -> RFQ-HOST-{{N:03d}}, BIOS N -> RFQ-BIOS-{{N:03d}}。若無固定格式則填 AI auto",
  "table_structure": {{
    "id_col": "ID 欄名稱",
    "requirement_col": "需求內容欄名稱",
    "comment_col": "備註欄名稱"
  }},
  "main_files": ["主需求檔案清單"],
  "appendix_files": ["附件/問卷檔案清單"],
  "language": "en/zh/mixed",
  "confidence": 0.0,
  "notes": ""
}}

注意：
- req_id_rule 只用 ASCII 字元，不要用 Unicode 箭頭符號
- table_structure 應以 is_main_requirement=true 且有明確需求清單的檔案為準（例如有 SRS Ref.# + Requirement 欄的 sheet，而非純規格對照表）
- 合約模板（Agreement Template）、報價單（Quote Template）不是 main_files
- req_id_rule 範例：
  - 若 req_id 欄是章節編號（4.1、4.2.1），填：SRS Ref.# directly（保留原始編號，不轉換）
  - 若 req_id 欄是 REQ-VS9-HIGH-SECU-001 這種格式，填：REQ-ID directly
  - 若 req_id 欄是 HOST N / BIOS N / BMC N，填：HOST N -> RFQ-HOST-{{N:03d}}, BIOS N -> RFQ-BIOS-{{N:03d}}, BMC N -> RFQ-BMC-{{N:03d}}
  - 若無 req_id 欄，填：AI auto"""

    try:
        resp = client.chat.completions.create(
            model=model, temperature=0,
            messages=[
                {"role": "system", "content": _SYSTEM_JSON},
                {"role": "user", "content": prompt2},
            ],
        )
        text = (resp.choices[0].message.content or "").strip()
        schema = parse_json_response(text, model=model)
        safe = str(schema.get("rfq_format", "")).encode("ascii", errors="replace").decode("ascii")
        print(f"[INFO] Final schema: format={safe} confidence={schema.get('confidence')}")
        return schema
    except Exception as e:
        print(f"[WARN] analyze_document_schema round2 failed: {str(e).encode('ascii', errors='replace').decode('ascii')[:300]}")
    return {}


def load_or_create_schema(client, model: str, rfq_dir: Path, meta_dir: Path) -> dict:
    """
    讀取已存在的 doc_schema.json，若不存在則呼叫 LLM 分析並儲存。
    """
    schema_path = meta_dir / "doc_schema.json"
    if schema_path.exists():
        try:
            schema = json.loads(schema_path.read_text(encoding="utf-8"))
            print(f"[INFO] Loaded existing doc_schema from {schema_path}")
            return schema
        except Exception:
            pass

    print("[INFO] Analyzing document schema with LLM...")
    schema = analyze_document_schema(client, model, rfq_dir)
    if schema:
        meta_dir.mkdir(parents=True, exist_ok=True)
        schema_path.write_text(json.dumps(schema, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[INFO] doc_schema saved to {schema_path}")
    return schema


def _parse_chunk(val):
    """Parse chunk identifier that may be int or string like 'ROW 12'."""
    if val is None:
        return 0
    m = re.search(r'\d+', str(val).strip())
    return int(m.group()) if m else 0


# -----------------------------
# Basic helpers
# -----------------------------
def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def read_txt(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")


def read_md(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")


# -----------------------------
# ID detection (vendor-agnostic-ish)
# -----------------------------
def detect_row_id(row_text: str) -> str:
    """
    Detect common requirement IDs across vendors.
    Returns normalized ID like HOST-36 / BIOS-4 / BMC-6 / REQ-123 / SR-4.3.21 ...
    """
    s = (row_text or "").strip()

    patterns = [
        r"\b(HOST|BIOS|BMC)\s*(\d+)\b",          # HOST 36
        r"\bREQ[-_ ]?(\d+)\b",                   # REQ-123
        r"\bSR[-_ ]?(\d+(\.\d+)+)\b",            # SR-4.3.21
        r"\bR[-_ ]?(\d+(\.\d+)+)\b",             # R-1.2.3
        r"\bSEC[-_ ]?(\d+)\b",                   # SEC-07
        r"\bITEM\s*(\d+)\b",                     # Item 12
    ]

    for pat in patterns:
        m = re.search(pat, s, flags=re.IGNORECASE)
        if not m:
            continue

        g0 = m.group(0).upper()
        if g0.startswith(("HOST", "BIOS", "BMC")):
            return f"{m.group(1).upper()}-{m.group(2)}"
        if g0.startswith("REQ"):
            return f"REQ-{m.group(1)}"
        if g0.startswith("SR"):
            return f"SR-{m.group(1)}"
        if g0.startswith("R"):
            return f"R-{m.group(1)}"
        if g0.startswith("SEC"):
            return f"SEC-{m.group(1)}"
        if g0.startswith("ITEM"):
            return f"ITEM-{m.group(1)}"

    return ""


# -----------------------------
# DOCX reading: paragraphs + tables
# -----------------------------
def read_docx_blocks(p: Path) -> List[str]:
    """
    Read DOCX blocks including:
    - normal paragraphs
    - table rows (each row becomes one block)
    """
    if docx is None:
        raise RuntimeError("python-docx not installed. pip install python-docx")

    d = docx.Document(str(p))
    blocks: List[str] = []

    # 1) paragraphs
    for para in d.paragraphs:
        t = (para.text or "").strip()
        if t:
            t = re.sub(r"\s+", " ", t).strip()
            blocks.append(t)

    # 2) tables: each row as a block (with IBM Matrix detection)
    for ti, table in enumerate(d.tables, start=1):
        # --- Detect IBM Matrix header structure ---
        ibm_schema: Dict[str, int] = {}  # col_name -> col_index
        if table.rows:
            header_cells = []
            for cell in table.rows[0].cells:
                ct = (cell.text or "").strip()
                ct = re.sub(r"\s+", " ", ct).strip()
                header_cells.append(ct)
            header_lower = [h.lower() for h in header_cells]
            # Check if this is an IBM Matrix table
            has_req_col = any("requirement" in h for h in header_lower)
            if has_req_col:
                for ci, h in enumerate(header_lower):
                    if "item" in h or h in ("no", "no.", "#", "id"):
                        ibm_schema["item"] = ci
                    elif "ibm requirement" in h or (h == "requirement"):
                        ibm_schema["requirement"] = ci
                    elif "ibm comment" in h or "comment" in h and "vendor" not in h:
                        ibm_schema["ibm_comment"] = ci
                    elif "vendor" in h:
                        ibm_schema["vendor_comment"] = ci
                # Must have at least the requirement column
                if "requirement" not in ibm_schema:
                    ibm_schema = {}

        use_ibm = bool(ibm_schema and "requirement" in ibm_schema)
        if use_ibm:
            print(f"[INFO] IBM Matrix detected in table {ti}: {ibm_schema}")

        if use_ibm:
            # --- IBM Matrix: skip header, process data rows ---
            for ri, row in enumerate(table.rows[1:], 2):
                raw_cells = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    ct = re.sub(r"\s+", " ", ct).strip()
                    raw_cells.append(ct)

                req_ci = ibm_schema["requirement"]
                req_text = raw_cells[req_ci] if req_ci < len(raw_cells) else ""
                if not req_text:
                    continue

                item = ""
                if "item" in ibm_schema:
                    item_ci = ibm_schema["item"]
                    item = raw_cells[item_ci] if item_ci < len(raw_cells) else ""

                comment = ""
                if "ibm_comment" in ibm_schema:
                    cmt_ci = ibm_schema["ibm_comment"]
                    comment = raw_cells[cmt_ci] if cmt_ci < len(raw_cells) else ""

                rid = item or detect_row_id(req_text)
                row_id_hint = f"[ROW_ID={rid}]" if rid else ""

                comment_hint = f" [IBM_COMMENT: {comment}]" if comment else ""
                blocks.append(f"[TABLE {ti} ROW {ri}]{row_id_hint} {req_text}{comment_hint}")
        else:
            # --- Default: merge all cells (skip header) ---
            for ri, row in enumerate(table.rows[1:], 2):
                raw_cells = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    ct = re.sub(r"\s+", " ", ct).strip()
                    raw_cells.append(ct)

                cells = [c for c in raw_cells if c]
                if not cells:
                    continue

                row_text = " | ".join(cells)
                rid = detect_row_id(row_text)
                row_id_hint = f" [ROW_ID={rid}]" if rid else ""
                blocks.append(f"[TABLE {ti} ROW {ri}]{row_id_hint} {row_text}")

    return blocks


# -----------------------------
# .doc reader (python-docx with plain-text fallback)
# -----------------------------
def read_doc_blocks(p: Path) -> List[str]:
    """
    Read .doc file: try python-docx first, fall back to plain text on failure.
    """
    if docx is not None:
        try:
            return read_docx_blocks(p)
        except Exception:
            pass
    return [p.read_text(encoding="utf-8", errors="ignore")]


# -----------------------------
# .xlsx / .xls reader
# -----------------------------
_RE_FILE_SHEET_ROW = re.compile(r"\[FILE:(.+?)\s+SHEET:(.+?)\s+ROW\s+(\d+)\]")
_RE_SHEET_ROW = re.compile(r"\[SHEET:(.+?)\s+ROW\s+(\d+)\]")


def _parse_sheet_rows(chunk: str) -> List[Tuple[str, str, int, str]]:
    """Parse chunk into list of (file, sheet_name, row_num, row_text) from block markers.

    Supports both [FILE:xxx SHEET:xxx ROW N] and legacy [SHEET:xxx ROW N] formats.
    """
    results: List[Tuple[str, str, int, str]] = []

    # Try new format first: [FILE:xxx SHEET:xxx ROW N]
    markers = list(_RE_FILE_SHEET_ROW.finditer(chunk))
    if markers:
        for idx, m in enumerate(markers):
            filename = m.group(1)
            sheet = m.group(2)
            row = int(m.group(3))
            start = m.end()
            next_m = markers[idx + 1] if idx + 1 < len(markers) else None
            row_text = chunk[start:next_m.start()].strip() if next_m else chunk[start:].strip()
            results.append((filename, sheet, row, row_text))
        return results

    # Fallback: legacy [SHEET:xxx ROW N]
    for m in _RE_SHEET_ROW.finditer(chunk):
        sheet = m.group(1)
        row = int(m.group(2))
        start = m.end()
        next_m = _RE_SHEET_ROW.search(chunk, start)
        row_text = chunk[start:next_m.start()].strip() if next_m else chunk[start:].strip()
        results.append(("", sheet, row, row_text))
    return results


def _match_sheet_row(req_text: str, parsed: List[Tuple[str, str, int, str]]) -> Tuple[str, str, int]:
    """Find best matching (file, sheet, row) for a requirement text. Returns ("", "", -1) if no match."""
    if not parsed:
        return "", "", -1
    req_lower = req_text.lower()
    best_overlap = 0
    best = parsed[0]  # fallback to first
    for filename, sheet, row, row_text in parsed:
        rt_lower = row_text.lower()
        # check if requirement is a substring of row or vice versa
        if req_lower in rt_lower or rt_lower in req_lower:
            return filename, sheet, row
        # simple word overlap
        req_words = set(req_lower.split())
        row_words = set(rt_lower.split())
        overlap = len(req_words & row_words)
        if overlap > best_overlap:
            best_overlap = overlap
            best = (filename, sheet, row, row_text)
    return best[0], best[1], best[2]


def detect_xlsx_schema(client, model: str, sheet_name: str, sample_rows: list) -> dict:
    """
    用 LLM 判斷 Excel sheet 的欄位結構。
    回傳 {"req_id_col": 0, "requirement_col": 1, "header_row": 2}
    若無法判斷則回傳 {}
    """
    if not sample_rows:
        return {}

    sample_text = ""
    for i, row in enumerate(sample_rows[:8]):
        cells = [str(c) if c is not None else "" for c in row]
        sample_text += f"Row {i+1}: {' | '.join(cells)}\n"

    prompt = f"""你是 Excel 欄位分析器。以下是 Excel sheet "{sheet_name}" 的前幾行資料：

{sample_text}

請判斷：
1. header_row: 哪一行是欄位標題（從 1 開始，若沒有標題則為 0）
2. req_id_col: 哪一欄是 requirement ID（從 0 開始，例如 REQ-001、SYS-1、4.1 等編號，若無則為 -1）
3. requirement_col: 哪一欄是需求內容文字（從 0 開始）

只輸出 JSON，例如：{{"header_row": 3, "req_id_col": 0, "requirement_col": 1}}
若無法判斷輸出：{{}}"""

    try:
        resp = client.chat.completions.create(
            model=model,
            temperature=0,
            messages=[
                {"role": "system", "content": _SYSTEM_JSON},
                {"role": "user", "content": prompt},
            ],
        )
        text = (resp.choices[0].message.content or "").strip()
        return parse_json_response(text, model=model)
    except Exception:
        pass
    return {}


def read_xlsx_blocks(p: Path, client=None, model: str = "gpt-4.1-mini") -> List[str]:
    """
    Read .xlsx / .xls using openpyxl.
    Each non-empty row becomes one block:
      [FILE:{filename} SHEET:{name} ROW {n}]{row_id_hint} cell1 | cell2 | ...
    If client is provided, uses LLM to detect column schema for structured extraction.
    """
    if openpyxl is None:
        raise RuntimeError("openpyxl not installed. pip install openpyxl")

    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    blocks: List[str] = []
    filename = p.name

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_rows = list(ws.iter_rows(values_only=True))

        # --- LLM schema detection (optional) ---
        schema: dict = {}
        if client is not None and all_rows:
            schema = detect_xlsx_schema(client, model, sheet_name, all_rows[:8])
            if schema:
                _schema_safe = str(schema).encode('ascii', errors='replace').decode('ascii')
                print(f"[INFO] xlsx schema detected for sheet '{sheet_name}': {_schema_safe}")

        header_row = schema.get("header_row", 0)
        req_id_col = schema.get("req_id_col", -1)
        req_col = schema.get("requirement_col", -1)
        use_schema = bool(schema and req_col >= 0)

        for ri, row in enumerate(all_rows, start=1):
            # Skip header row(s) when schema is detected
            if use_schema and header_row > 0 and ri <= header_row:
                continue

            if use_schema:
                # Structured extraction using detected columns
                req_text = str(row[req_col]).strip() if req_col < len(row) and row[req_col] is not None else ""
                if not req_text:
                    continue
                rid = ""
                if req_id_col >= 0 and req_id_col < len(row) and row[req_id_col] is not None:
                    rid = str(row[req_id_col]).strip()
                if not rid:
                    rid = detect_row_id(req_text)
                row_id_hint = f" [ROW_ID={rid}]" if rid else ""
                blocks.append(f"[FILE:{filename} SHEET:{sheet_name} ROW {ri}]{row_id_hint} {req_text}")
            else:
                # Fallback: merge all cells (original behaviour)
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if not cells:
                    continue
                row_text = " | ".join(cells)
                rid = detect_row_id(row_text)
                row_id_hint = f" [ROW_ID={rid}]" if rid else ""
                blocks.append(f"[FILE:{filename} SHEET:{sheet_name} ROW {ri}]{row_id_hint} {row_text}")

    wb.close()
    return blocks


# -----------------------------
# .pdf reader
# -----------------------------
def read_pdf_blocks(p: Path) -> List[str]:
    """
    Read .pdf using pdfplumber. Each page becomes one block:
      [PDF PAGE {n}] {text}
    """
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not installed. pip install pdfplumber")

    blocks: List[str] = []
    with pdfplumber.open(str(p)) as pdf:
        for n, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                text = re.sub(r"\s+", " ", text).strip()
                blocks.append(f"[PDF PAGE {n}] {text}")
    return blocks


# -----------------------------
# Direct spreadsheet parser (no LLM)
# -----------------------------
def extract_simple_list_xlsx(fp: Path, doc_schema: dict) -> List[Dict[str, Any]]:
    """Direct extraction for simple_list format xlsx with identifiable columns."""
    if openpyxl is None:
        raise RuntimeError("openpyxl not installed. pip install openpyxl")

    table_struct = doc_schema.get("table_structure", {})
    id_col_name = table_struct.get("id_col", "ID")
    req_col_name = table_struct.get("requirement_col", "Question")
    ans_col_name = table_struct.get("comment_col", "Answer")

    wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
    reqs: List[Dict[str, Any]] = []

    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        header = [str(c or "").strip() for c in rows[0]]
        id_idx = next((i for i, h in enumerate(header) if h == id_col_name), None)
        req_idx = next((i for i, h in enumerate(header) if h == req_col_name), None)
        ans_idx = next((i for i, h in enumerate(header) if h == ans_col_name), None)

        if req_idx is None:
            print(f"[WARN] Sheet '{sn}': column '{req_col_name}' not found in header {header}, skipping")
            continue

        print(f"[INFO] Direct parse: sheet='{sn}' id_col={id_idx} req_col={req_idx} ans_col={ans_idx} rows={len(rows)-1}")

        for ri, row in enumerate(rows[1:], start=1):
            cells = list(row)
            req_text = str(cells[req_idx] or "").strip() if req_idx < len(cells) else ""
            if not req_text:
                continue
            raw_id = str(cells[id_idx] or "").strip() if id_idx is not None and id_idx < len(cells) else ""
            answer = str(cells[ans_idx] or "").strip() if ans_idx is not None and ans_idx < len(cells) else ""
            req_id = raw_id if raw_id else f"AUTO-{fp.name}-{ri}"
            reqs.append({
                "req_id": req_id,
                "requirement": req_text,
                "notes": answer,
                "confidence": 1.0,
                "source": {"file": fp.name, "sheet": sn, "row": ri + 1},
            })

    wb.close()
    print(f"[OK] Direct parse: {fp.name} -> {len(reqs)} requirements")
    return reqs


# ── Spec-category mapping for relaxed extraction ─────────────────────────────
_SPEC_CATEGORIES = {
    "CPU":           ["cpu", "processor", "soc", "marvell", "amd", "intel", "arm"],
    "Memory":        ["dram", "memory", "ram", "ddr"],
    "Storage":       ["emmc", "ssd", "nvme", "storage", "flash", "boot"],
    "Network":       ["ethernet", "switch", "fpga", "sfp", "rj45", "nic", "management"],
    "Wireless":      ["wifi", "wi-fi", "lte", "5g", "wwan", "antenna", "sim", "carrier",
                      "triband", "ieee802"],
    "Display":       ["display", "lcd", "led", "screen", "console"],
    "Power":         ["power", "poe", "battery", "dc", "ac", "fanless", "voltage"],
    "Certification": ["cert", "fips", "ptcrb", "tecc", "gcf", "fcc", "ul ", "ce mark"],
    "Mechanical":    ["wallmount", "rackmount", "pcba", "pcie", "usb", "m.2"],
    "Thermal":       ["temperature", "thermal", "operating temp"],
    "Performance":   ["performance", "throughput", "gbps", "mbps"],
}


def _guess_spec_category(label: str) -> str:
    """Map a spec row label to a canonical category."""
    ll = (label or "").lower()
    for cat, keywords in _SPEC_CATEGORIES.items():
        for kw in keywords:
            if kw in ll:
                return cat
    return "General"


def extract_spec_reference_xlsx(fp: Path, doc_schema: dict) -> List[Dict[str, Any]]:
    """Relaxed extraction for spec_reference xlsx: each row with a label + values becomes
    a derived requirement. No LLM needed."""
    if openpyxl is None:
        raise RuntimeError("openpyxl not installed. pip install openpyxl")

    wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
    reqs: List[Dict[str, Any]] = []

    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Detect header row (first row with multiple non-empty cells is likely SKU headers)
        header = [str(c or "").strip() for c in rows[0]]
        sku_names = [h for h in header[1:] if h]  # columns after label

        print(f"[INFO] Spec-reference parse: sheet='{sn}' SKUs={sku_names[:5]} rows={len(rows)}")

        for ri, row in enumerate(rows[1:] if sku_names else rows, start=1):
            cells = [str(c or "").strip() if c is not None else "" for c in row]
            if not cells:
                continue

            label = cells[0].strip()
            if not label:
                continue

            # Collect non-empty values from SKU columns
            values = [cells[j].strip() for j in range(1, len(cells)) if j < len(cells) and cells[j].strip()]
            if not values and len(cells) == 1:
                # Single-column label with no values — skip (section header)
                continue

            # Build requirement text
            if sku_names and values:
                pairs = []
                for j, val in enumerate(cells[1:], start=0):
                    v = str(val or "").strip()
                    if v and j < len(sku_names):
                        pairs.append(f"{sku_names[j]}: {v}")
                    elif v:
                        pairs.append(v)
                req_text = f"{label}: {', '.join(pairs)}" if pairs else label
            elif values:
                req_text = f"{label}: {', '.join(values)}"
            else:
                req_text = label

            spec_cat = _guess_spec_category(label)
            reqs.append({
                "req_id": f"AUTO-{fp.name}-{ri}",
                "requirement": req_text,
                "notes": "Derived from spec table — no explicit shall/must",
                "confidence": 0.5,
                "derived_requirement": True,
                "spec_category": spec_cat,
                "source": {"file": fp.name, "sheet": sn, "row": ri + 1},
            })

    wb.close()
    print(f"[OK] Spec-reference parse: {fp.name} -> {len(reqs)} derived requirements")
    return reqs


# -----------------------------
# ── Checklist / compliance table parser ───────────────────────────────────────

_CHECKLIST_LABEL_KEYWORDS = re.compile(
    r"requirement|description|model|specification|item|feature", re.IGNORECASE
)
_CHECKLIST_COMPLY_KEYWORDS = re.compile(
    r"comply|compliance|compliant", re.IGNORECASE
)
_CHECKLIST_REF_KEYWORDS = re.compile(
    r"ref|id|#|srs|section", re.IGNORECASE
)
_CHECKLIST_PRIORITY_KEYWORDS = re.compile(
    r"priority|level|must|importance", re.IGNORECASE
)

# Section headers: short titles without specific content (skip these)
_SECTION_HEADER_RE = re.compile(
    r"^[\d\.]+\s*$"  # just a number like "4.1"
)


def _detect_checklist_header(rows: List[tuple], max_scan: int = 10) -> Optional[Dict[str, Any]]:
    """Scan first max_scan rows to find a checklist header.

    Returns dict with header_row index and column indices, or None.
    """
    for i, row in enumerate(rows[:max_scan]):
        cells = [str(c or "").strip() for c in row]
        cells_lower = [c.lower() for c in cells]
        joined = " ".join(cells_lower)

        has_label = bool(_CHECKLIST_LABEL_KEYWORDS.search(joined))
        has_comply = bool(_CHECKLIST_COMPLY_KEYWORDS.search(joined))

        if not (has_label and has_comply):
            continue

        # Find column indices
        ref_idx = None
        req_idx = None
        priority_idx = None
        comply_idx = None
        note_idx = None

        for ci, cl in enumerate(cells_lower):
            if ref_idx is None and _CHECKLIST_REF_KEYWORDS.search(cl):
                ref_idx = ci
            elif req_idx is None and _CHECKLIST_LABEL_KEYWORDS.search(cl):
                req_idx = ci
            elif priority_idx is None and _CHECKLIST_PRIORITY_KEYWORDS.search(cl):
                priority_idx = ci
            elif comply_idx is None and _CHECKLIST_COMPLY_KEYWORDS.search(cl):
                comply_idx = ci

        # Note/clarification column: look for "partially" or "comment" or "note"
        for ci, cl in enumerate(cells_lower):
            if any(kw in cl for kw in ("partial", "clarif", "comment", "note", "remark")):
                if ci != comply_idx:
                    note_idx = ci
                    break

        if req_idx is not None:
            return {
                "header_row": i,
                "ref_idx": ref_idx,
                "req_idx": req_idx,
                "priority_idx": priority_idx,
                "comply_idx": comply_idx,
                "note_idx": note_idx,
                "header_cells": cells,
            }
    return None


def _map_priority(raw: str) -> str:
    """Map checklist priority codes to must_level values."""
    p = (raw or "").strip().upper()
    if p in ("M", "MUST", "MANDATORY", "H", "HIGH"):
        return "MUST"
    if p in ("S", "SHOULD", "MEDIUM"):
        return "SHOULD"
    if p in ("L", "LOW", "MAY", "O", "OPTIONAL"):
        return "MAY"
    return "INFO"


def _is_section_header(req_text: str, ref_text: str, priority_raw: str) -> bool:
    """True if this row is a section header, not an actual requirement.

    In a compliance checklist, rows WITH a priority value (M, H, L) are real items,
    even if the text is short. Only skip rows that are both short AND have no priority.
    """
    t = (req_text or "").strip()
    # If priority is set, this is a real checklist item regardless of length
    if (priority_raw or "").strip():
        return False
    word_count = len(re.findall(r"\w+", t))
    if word_count <= 2 and not re.search(r"\b(shall|must|required|support|provide)\b", t, re.IGNORECASE):
        return True
    return False


def extract_checklist_xlsx(fp: Path) -> List[Dict[str, Any]]:
    """Parse compliance checklist xlsx. Auto-detects header per sheet.
    Returns requirement items with derived_requirement=False."""
    if openpyxl is None:
        raise RuntimeError("openpyxl not installed. pip install openpyxl")

    wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
    reqs: List[Dict[str, Any]] = []
    sheets_parsed = 0

    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        header_info = _detect_checklist_header(rows)
        if header_info is None:
            print(f"[INFO] Checklist: sheet '{sn}' — no checklist header, skipping")
            continue

        hi = header_info["header_row"]
        ref_idx = header_info["ref_idx"]
        req_idx = header_info["req_idx"]
        pri_idx = header_info["priority_idx"]
        comply_idx = header_info["comply_idx"]
        note_idx = header_info["note_idx"]

        print(f"[INFO] Checklist parse: sheet='{sn}' header=R{hi} "
              f"ref={ref_idx} req={req_idx} pri={pri_idx} comply={comply_idx} note={note_idx} "
              f"data_rows={len(rows)-hi-1}")

        sheets_parsed += 1
        skipped_headers = 0

        for ri, row in enumerate(rows[hi + 1:], start=hi + 2):
            cells = [str(c or "").strip() if c is not None else "" for c in row]

            req_text = cells[req_idx] if req_idx < len(cells) else ""
            if not req_text.strip():
                continue

            ref_text = cells[ref_idx] if ref_idx is not None and ref_idx < len(cells) else ""
            priority_raw = cells[pri_idx] if pri_idx is not None and pri_idx < len(cells) else ""

            if _is_section_header(req_text, ref_text, priority_raw):
                skipped_headers += 1
                continue
            comply_val = cells[comply_idx] if comply_idx is not None and comply_idx < len(cells) else ""
            note_val = cells[note_idx] if note_idx is not None and note_idx < len(cells) else ""

            # Build notes from comply status + clarification
            notes_parts = []
            if comply_val:
                notes_parts.append(f"Comply: {comply_val}")
            if note_val:
                notes_parts.append(note_val)

            req_id = ref_text if ref_text else f"AUTO-{fp.name}-{sn}-{ri}"

            reqs.append({
                "req_id": req_id,
                "requirement": req_text,
                "notes": " | ".join(notes_parts) if notes_parts else "",
                "confidence": 1.0,
                "derived_requirement": False,
                "must_level": _map_priority(priority_raw),
                "source": {"file": fp.name, "sheet": sn, "row": ri},
            })

        if skipped_headers:
            print(f"[INFO] Checklist: sheet '{sn}' — skipped {skipped_headers} section headers")

    wb.close()
    print(f"[OK] Checklist parse: {fp.name} -> {len(reqs)} requirements from {sheets_parsed} sheets")
    return reqs


def has_checklist_sheets(fp: Path) -> bool:
    """Quick check: does this xlsx have any sheet with a checklist header?"""
    if openpyxl is None:
        return False
    try:
        wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
        for sn in wb.sheetnames:
            ws = wb[sn]
            rows = list(ws.iter_rows(values_only=True, max_row=10))
            if _detect_checklist_header(rows) is not None:
                wb.close()
                return True
        wb.close()
    except Exception:
        pass
    return False


# -----------------------------
# Chunking
# -----------------------------
def force_split_by_length(text: str, max_chars: int) -> List[str]:
    t = (text or "").strip()
    if not t:
        return []
    if len(t) <= max_chars:
        return [t]

    sentences = re.split(r"(?<=[\.\!\?\；\。\!\?])\s+", t)
    chunks: List[str] = []
    buf = ""
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        if len(buf) + len(s) + 1 <= max_chars:
            buf = (buf + " " + s).strip()
        else:
            if buf:
                chunks.append(buf)
            if len(s) > max_chars:
                for i in range(0, len(s), max_chars):
                    part = s[i:i + max_chars].strip()
                    if part:
                        chunks.append(part)
                buf = ""
            else:
                buf = s
    if buf:
        chunks.append(buf)
    return chunks


def split_chunks_generic(text: str, max_chars: int = 3500) -> List[str]:
    t = re.sub(r"\r\n", "\n", text or "")
    blocks = [b.strip() for b in re.split(r"\n\s*\n", t) if b.strip()]
    chunks: List[str] = []
    for b in blocks:
        chunks.extend(force_split_by_length(b, max_chars))
    if not chunks and (text or "").strip():
        chunks = force_split_by_length((text or "").strip(), max_chars)
    return chunks


def chunks_from_blocks(blocks: List[str], max_chars: int = 1200, group_size: int = 4) -> List[str]:
    chunks: List[str] = []
    buf: List[str] = []
    buf_len = 0

    def flush():
        nonlocal buf, buf_len
        if buf:
            chunks.append("\n".join(buf).strip())
            buf = []
            buf_len = 0

    for b in blocks:
        if len(b) > max_chars:
            flush()
            # Table rows with ROW_ID must stay intact — never split them
            if re.search(r'\[ROW_ID=', b):
                chunks.append(b)
            else:
                for part in force_split_by_length(b, max_chars):
                    chunks.append(part)
            continue

        if buf and (len(buf) >= group_size or buf_len + len(b) + 2 > max_chars):
            flush()
        buf.append(b)
        buf_len += len(b) + 2

    flush()
    return chunks


# -----------------------------
# Prompt
# -----------------------------
def build_prompt(file_name: str, chunk_index: int, chunk_text: str, doc_schema: dict = None) -> str:
    schema_hint = ""
    if doc_schema:
        rule = doc_schema.get("req_id_rule", "")
        fmt = doc_schema.get("rfq_format", "")
        if rule and rule != "AI 自動編號":
            schema_hint = f"""
【本份 RFQ 的 req_id 規則（優先遵守）】
文件格式：{fmt}
req_id 轉換規則：{rule}
例如：若看到 HOST 1 或 [ROW_ID=HOST-1]，req_id 必須設為 RFQ-HOST-001
"""

    return f"""
你是 RFQ/RFI 文件的「Requirements 抽取器」。

任務：從下面這段文件內容中，抽取可驗證、可回覆的 requirement 條目，輸出成 JSON。
{schema_hint}
重點原則：
- 自由段落（非表格）：請拆細，一句一個 requirement（尤其含 shall/must/required/need/必須/需要/不得）
- 表格列（含 [ROW_ID=...] 或 [TABLE N ROW N] 標記）：整列作為一個 requirement，不要拆分 cell 內容。
  保留整格原文，包含 | 分隔的多個項目。

輸出規則（非常重要）：
- 只輸出 JSON（不要任何解釋、不要 markdown）
- 每個 chunk 最多輸出 10 條 requirement。若內容超過 10 條，優先保留含 shall/must/required/obligation/deliverable 的條目，其餘省略。

- req_id 規則（最高優先順序）：
  - 若文字包含 [ROW_ID=...]，req_id 必須使用該值（例如 HOST-36 / BIOS-4 / BMC-6）
  - 若沒有 [ROW_ID=...]，req_id 用自動格式 AUTO-<file>-<chunk>-<seq>

- 表格列不拆分原則：
  - 有 [ROW_ID=...] 或 [TABLE N ROW N] 標記的列，整列只產生一個 requirement 條目
  - requirement 欄位保留整格原文（包含 | 分隔的所有內容）
  - 不要把一個 table row 拆成多個 requirement

- 表格列處理（只在「可辨識欄位」時才啟用）：
  - 若文字包含 [IBM_MATRIX]，或出現明顯欄位結構（例如含有多個 '|' 分隔欄位），請把該列拆解後，只抽取「Requirement」欄位的內容作為 requirement。
    * 若該列出現 'IBM Requirement' 欄位，優先取 IBM Requirement 欄位。
    * 若沒有 IBM Requirement，但有 'Requirement' 欄位，取 Requirement 欄位。
    * Comment/Notes/Vendor Comment/Response 等欄位不要併入 requirement（可放到 notes）。
  - 若不符合以上條件（例如一般段落、一般條列、或表格但看不出欄位意義），請不要假設欄位順序，直接依句意抽 requirement。

- 雜訊/定義句（避免污染）：
  - 若句子像名詞定義或縮寫解釋（例如「CRU: Customer Replaceable Unit」），仍可輸出，但請：
    must_level = "INFO"
    notes = "GLOSSARY/DEFINITION"
    confidence 建議較低（例如 0.4~0.7）

- JSON 格式如下：
{{
  "requirements": [
    {{
      "req_id": "AUTO-{file_name}-{chunk_index}-001",
      "requirement": "一句清楚的要求（英文原文可保留）",
      "source": {{
        "file": "{file_name}",
        "chunk": {chunk_index}
      }},
      "notes": "",
      "confidence": 0.0
    }}
  ]
}}
- confidence: 0.0~1.0
- 若此段沒有可抽取的 requirement，輸出：{{"requirements": []}}

開始內容：
----------------
{chunk_text}
----------------
""".strip()


# -----------------------------
# LLM call with retry
# -----------------------------
_SYSTEM_JSON = (
    "You output STRICT JSON only. "
    "Return exactly one JSON object. "
    "Do not wrap in markdown fences. "
    "Do not add any text before or after the JSON."
)


def call_llm_json_with_retry(client: OpenAI, model: str, prompt: str, retries: int = 3) -> Dict[str, Any]:
    last_err = None
    for attempt in range(1, retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                temperature=0.1,
                max_tokens=4096,
                messages=[
                    {"role": "system", "content": _SYSTEM_JSON},
                    {"role": "user", "content": prompt},
                ],
            )
            text = (resp.choices[0].message.content or "").strip()
            finish = getattr(resp.choices[0], "finish_reason", None)
            if finish == "length":
                print(f"[WARN] LLM output truncated (finish_reason=length), chunk may be incomplete")
            return parse_json_response(text, model=model)
        except Exception as e:
            last_err = e
            wait_s = min(2 ** attempt, 10)
            err_safe = str(e).encode("ascii", errors="replace").decode("ascii")[:300]
            print(f"[WARN] LLM call failed (attempt {attempt}/{retries}): {err_safe} -> sleep {wait_s}s")
            time.sleep(wait_s)

    err_safe = str(last_err).encode("ascii", errors="replace").decode("ascii")[:300]
    raise RuntimeError(f"LLM call failed after {retries} retries: {err_safe}")


# -----------------------------
# Resume / partial jsonl
# -----------------------------
def load_done_keys(partial_path: Path,
                   include_failed_as_done: bool = True
                   ) -> Tuple[Set[Tuple[str, int]], List[Dict[str, Any]]]:
    """
    Read jsonl and return:
    - done keys: {(file, chunk)}
    - flattened requirements list (dedup later)

    Phase 4.6G: records carrying failed_chunk=true are normally treated as
    done (so resume skips them like any other completed chunk). Pass
    include_failed_as_done=False to make resume re-attempt them — the caller
    wires this to --retry-failed-chunks.
    """
    done: Set[Tuple[str, int]] = set()
    reqs_all: List[Dict[str, Any]] = []
    if not partial_path.exists():
        return done, reqs_all

    for line in partial_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
        except Exception:
            continue

        f = obj.get("file")
        c = obj.get("chunk")
        is_failed = bool(obj.get("failed_chunk", False))
        if f and isinstance(c, int):
            if not (is_failed and not include_failed_as_done):
                done.add((f, c))

        rs = obj.get("requirements") or []
        if isinstance(rs, list):
            reqs_all.extend(rs)

    return done, reqs_all


def append_partial(partial_path: Path, file_name: str, chunk_index: int,
                   requirements: List[Dict[str, Any]],
                   failed: bool = False) -> None:
    """Append one chunk's outcome to partial.jsonl. Phase 4.6G adds the
    `failed` flag: when True, the record carries `failed_chunk=true` so
    load_done_keys can tell a legit empty-requirements chunk apart from a
    chunk we gave up on after retry exhaust."""
    partial_path.parent.mkdir(parents=True, exist_ok=True)
    rec = {"file": file_name, "chunk": chunk_index, "requirements": requirements}
    if failed:
        rec["failed_chunk"] = True
    with partial_path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")


def _append_extract_error(errors_path: Path, file_name: str, chunk_index: int,
                          total: int, error: str, model: str,
                          chunk_chars: int) -> None:
    """Phase 4.6G — append a record describing a chunk we gave up on after
    retry exhaust. Each line is one independent failure event; the file is
    append-only across runs so PMs can see history."""
    errors_path.parent.mkdir(parents=True, exist_ok=True)
    rec = {
        "file":        file_name,
        "chunk":       chunk_index,
        "total":       total,
        "error":       error,
        "model":       model,
        "ts":          now_iso(),
        "chunk_chars": chunk_chars,
    }
    with errors_path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")


def dedup_requirements(reqs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Deduplicate by (req_id, source.file, source.chunk, requirement).
    This prevents duplicates when resume + rerun edge cases.
    """
    seen = set()
    out = []
    for r in reqs:
        src = r.get("source") or {}
        key = (
            (r.get("req_id") or "").strip(),
            (src.get("file") or "").strip(),
            _parse_chunk(src.get("chunk")),
            (r.get("requirement") or "").strip(),
        )
        if key in seen:
            continue
        seen.add(key)
        out.append(r)
    return out


# -----------------------------
# Main
# -----------------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--case", required=True, help=r"Case folder, e.g. inbound\20260129_IBM_RFQ")
    ap.add_argument("--runs", required=True, help=r"Runs root, e.g. runs")
    ap.add_argument("--model", default=None, help="Override model (default: resolved by LLM_PROVIDER / OPENAI_MODEL / INTERNAL_LLM_MODEL)")
    ap.add_argument("--max-chars", type=int, default=1200, help="Chunk size by characters")
    ap.add_argument("--group-size", type=int, default=4, help="DOCX: blocks per chunk")
    ap.add_argument("--resume", action="store_true", help="Resume from partial jsonl if exists")
    ap.add_argument("--reset-partial", action="store_true", help="Delete partial jsonl and rerun cleanly")
    ap.add_argument("--retries", type=int, default=3, help="LLM retries for each chunk")
    # Phase 4.6G — bad chunk soft-fail + threshold gate
    ap.add_argument("--max-failed-chunks", type=int, default=20,
                    help="Abort extract if absolute count of failed chunks in this run "
                         "exceeds this. 0 = abort on the first failure.")
    ap.add_argument("--max-failed-pct", type=float, default=5.0,
                    help="Abort extract if failed_chunks / attempted_chunks in this run "
                         "exceeds this percentage. attempted_chunks excludes resume-skipped.")
    ap.add_argument("--min-attempted-for-pct", type=int, default=50,
                    help="Disable --max-failed-pct gate until this many chunks attempted "
                         "in this run. Avoids tripping on small-sample noise.")
    ap.add_argument("--retry-failed-chunks", action="store_true",
                    help="On resume, re-attempt chunks previously marked failed_chunk=true. "
                         "Default: skip them like any other completed chunk.")
    args = ap.parse_args()

    if not is_available():
        raise RuntimeError(
            "LLM not configured. Set OPENAI_API_KEY (provider=openai) or "
            "LLM_PROVIDER=internal with INTERNAL_LLM_BASE_URL/INTERNAL_LLM_API_KEY/INTERNAL_LLM_MODEL."
        )

    client = get_client()
    args.model = args.model or get_model()

    base_dir = Path(__file__).resolve().parent
    case_dir = (base_dir / args.case).resolve() if not Path(args.case).is_absolute() else Path(args.case)
    rfq_dir = case_dir / "rfq"
    meta_case = case_dir / "meta" / "case.yaml"

    if not rfq_dir.exists():
        raise FileNotFoundError(f"rfq folder not found: {rfq_dir}")

    case_id = case_dir.name
    if meta_case.exists():
        try:
            import yaml
            m = yaml.safe_load(meta_case.read_text(encoding="utf-8")) or {}
            case_id = m.get("case_id") or case_id
        except Exception:
            pass

    runs_root = (base_dir / args.runs).resolve() if not Path(args.runs).is_absolute() else Path(args.runs)
    out_dir = runs_root / case_id
    out_dir.mkdir(parents=True, exist_ok=True)

    partial_path = out_dir / "requirements.partial.jsonl"

    if args.reset_partial and partial_path.exists():
        partial_path.unlink()
        print(f"[OK] Deleted partial: {partial_path}")

    # Phase 7: load Step 1.5 exclusion set early so we can also scrub any stale
    # partial.jsonl entries that were written before the file was excluded.
    excluded_files = load_excluded(case_dir)
    if excluded_files:
        print(f"[INFO] Step 1.5: {len(excluded_files)} file(s) marked excluded: "
              f"{sorted(excluded_files)}")

    done_keys: Set[Tuple[str, int]] = set()
    all_reqs: List[Dict[str, Any]] = []

    # Phase 4.6G — bad chunk soft-fail bookkeeping
    failed_chunks: List[Dict[str, Any]] = []
    attempted_chunks: int = 0
    errors_path: Path = out_dir / "extract_errors.jsonl"

    if args.resume and partial_path.exists():
        done_keys, partial_reqs = load_done_keys(
            partial_path,
            include_failed_as_done=not args.retry_failed_chunks,
        )
        if excluded_files:
            _orig_partial = len(partial_reqs)
            partial_reqs = [
                r for r in partial_reqs
                if str((r.get("source") or {}).get("file", "")) not in excluded_files
            ]
            done_keys = {k for k in done_keys if k[0] not in excluded_files}
            _dropped_partial = _orig_partial - len(partial_reqs)
            if _dropped_partial > 0:
                print(f"[INFO] Step 1.5: dropped {_dropped_partial} stale partial req(s) "
                      "from now-excluded file(s)")
        all_reqs.extend(partial_reqs)
        print(f"[OK] Resume enabled. Loaded done chunks: {len(done_keys)}; partial reqs: {len(partial_reqs)}")

    # 分析文件格式 schema（client 必然存在，因為上面已檢查 is_available）
    doc_schema: dict = {}
    if client:
        meta_dir = case_dir / "meta"
        doc_schema = load_or_create_schema(client, args.model, rfq_dir, meta_dir)
        if doc_schema:
            confidence = doc_schema.get("confidence", 0)
            fmt = doc_schema.get("rfq_format", "unknown")
            rule = doc_schema.get("req_id_rule", "")
            rule_safe = (rule or "").encode("ascii", errors="replace").decode("ascii")
            print(f"[INFO] Schema: format={fmt}, confidence={confidence}, req_id_rule={rule_safe}")
            if confidence < 0.7:
                print(f"[WARN] Low confidence schema ({confidence}). Please review inbound/meta/doc_schema.json")
        else:
            print("[WARN] Could not analyze document schema. Using default extraction.")

    def _get_file_schema(doc_schema: dict, filename: str) -> dict:
        """從 doc_schema 找到對應檔案的 schema"""
        if not doc_schema:
            return {}
        # 新格式：有 files list
        files_list = doc_schema.get("files", [])
        for f in files_list:
            if f.get("file", "") == filename or f.get("file", "").startswith(filename[:20]):
                return f
        # 舊格式：用 main_files / appendix_files 推斷 role
        result = dict(doc_schema)
        appendix = doc_schema.get("appendix_files") or []
        main = doc_schema.get("main_files") or []
        if any(filename == af or filename.startswith(af[:20]) for af in appendix):
            result["role"] = "appendix"
        elif any(filename == mf or filename.startswith(mf[:20]) for mf in main):
            result["role"] = "main_requirement"
        return result

    files: List[Path] = []
    for ext in ("*.docx", "*.doc", "*.xlsx", "*.xls", "*.pdf", "*.md", "*.txt"):
        files.extend(sorted(rfq_dir.glob(ext)))
    if not files:
        raise FileNotFoundError(f"No supported files found in: {rfq_dir} (docx/doc/xlsx/xls/pdf/md/txt)")

    # Phase 7: filter out files the PM marked Include=False in Step 1.5
    if excluded_files:
        _dropped_files = [fp.name for fp in files if fp.name in excluded_files]
        files = [fp for fp in files if fp.name not in excluded_files]
        if _dropped_files:
            print(f"[INFO] Step 1.5: skipping {len(_dropped_files)} excluded file(s): {_dropped_files}")
        if not files:
            raise FileNotFoundError(
                f"All input files in {rfq_dir} were marked Include=False in Step 1.5. "
                f"Re-enable some files via the UI (Step 1.5 → Save Selection) or "
                f"edit inbound/<case>/meta/file_selection.json directly."
            )

    for fp in files:
        name = fp.name
        suffix = fp.suffix.lower()

        file_schema = _get_file_schema(doc_schema, fp.name)
        file_role = file_schema.get("role", "main_requirement")
        file_req_id_rule = file_schema.get("req_id_rule", "AI auto")
        safe_rule = file_req_id_rule.encode("ascii", errors="replace").decode("ascii")
        print(f"[INFO] File: {fp.name} | role={file_role} | req_id_rule={safe_rule}")

        is_spec_ref = doc_schema.get("rfq_format") == "spec_reference"

        if file_role == "appendix" and not is_spec_ref:
            # Before skipping, check if this appendix xlsx is a compliance checklist
            if suffix in (".xlsx", ".xls") and has_checklist_sheets(fp):
                print(f"[INFO] Appendix {fp.name} has checklist sheets — parsing as checklist")
                checklist_reqs = extract_checklist_xlsx(fp)
                all_reqs.extend(checklist_reqs)
                append_partial(partial_path, name, 0, checklist_reqs)
                continue
            print(f"[SKIP] {fp.name} -- role=appendix, skipping")
            continue

        # ── Relaxed parse for spec_reference xlsx ──
        if is_spec_ref and suffix in (".xlsx", ".xls"):
            print(f"[INFO] spec_reference detected — using relaxed xlsx parser (no LLM)")
            derived_reqs = extract_spec_reference_xlsx(fp, doc_schema)
            all_reqs.extend(derived_reqs)
            append_partial(partial_path, name, 0, derived_reqs)
            continue

        if is_spec_ref and file_role == "appendix":
            print(f"[SKIP] {fp.name} -- spec_reference + non-xlsx appendix, skipping")
            continue

        # ── Direct parse for simple_list xlsx ──
        is_simple_list = (
            doc_schema.get("rfq_format") == "simple_list"
            and suffix in (".xlsx", ".xls")
            and doc_schema.get("table_structure", {}).get("requirement_col")
        )
        if is_simple_list:
            print(f"[INFO] simple_list detected — using direct xlsx parser (no LLM)")
            direct_reqs = extract_simple_list_xlsx(fp, doc_schema)
            all_reqs.extend(direct_reqs)
            append_partial(partial_path, name, 0, direct_reqs)
            continue

        # ── Standard LLM extraction path ──
        if suffix == ".docx":
            blocks = read_docx_blocks(fp)
            chunks = chunks_from_blocks(blocks, max_chars=args.max_chars, group_size=args.group_size)
        elif suffix == ".doc":
            blocks = read_doc_blocks(fp)
            chunks = chunks_from_blocks(blocks, max_chars=args.max_chars, group_size=args.group_size)
        elif suffix in (".xlsx", ".xls"):
            blocks = read_xlsx_blocks(fp, client=client, model=args.model)
            chunks = chunks_from_blocks(blocks, max_chars=args.max_chars, group_size=args.group_size)
        elif suffix == ".pdf":
            blocks = read_pdf_blocks(fp)
            chunks = chunks_from_blocks(blocks, max_chars=args.max_chars, group_size=args.group_size)
        else:
            text = read_md(fp) if suffix == ".md" else read_txt(fp)
            chunks = split_chunks_generic(text, max_chars=args.max_chars)

        print(f"[INFO] {name}: {len(chunks)} chunks")

        for i, chunk in enumerate(chunks, start=1):
            if args.resume and (name, i) in done_keys:
                print(f"[SKIP] {name} chunk {i}/{len(chunks)} already done")
                continue

            print(f"[PROGRESS] {name} chunk {i}/{len(chunks)}")
            attempted_chunks += 1
            prompt = build_prompt(name, i, chunk, doc_schema=file_schema or doc_schema)

            try:
                data = call_llm_json_with_retry(client, args.model, prompt, retries=args.retries)
            except RuntimeError as _e:
                # Phase 4.6G — bad chunk soft-fail: record, mark, continue.
                err_short = str(_e).encode("ascii", errors="replace").decode("ascii")[:500]
                print(f"[WARN] Skipping {name} chunk {i}/{len(chunks)} after retry exhaust: "
                      f"{err_short[:120]}")
                _append_extract_error(errors_path, name, i, len(chunks),
                                      err_short, args.model, len(chunk))
                append_partial(partial_path, name, i, [], failed=True)
                failed_chunks.append({"file": name, "chunk": i})

                # Threshold gate — abort if this run's failures look systemic.
                fail_pct = 100.0 * len(failed_chunks) / max(attempted_chunks, 1)
                abort_by_abs = len(failed_chunks) > args.max_failed_chunks
                abort_by_pct = (attempted_chunks >= args.min_attempted_for_pct
                                and fail_pct > args.max_failed_pct)
                if abort_by_abs or abort_by_pct:
                    raise RuntimeError(
                        f"Aborting extract: {len(failed_chunks)} failed chunk(s) "
                        f"({fail_pct:.1f}% of {attempted_chunks} attempted) "
                        f"exceeds threshold "
                        f"(max={args.max_failed_chunks}, max_pct={args.max_failed_pct}%, "
                        f"min_attempted_for_pct={args.min_attempted_for_pct}). "
                        f"See {errors_path} for details."
                    )
                continue

            reqs = data.get("requirements", []) or []

            # ---- Fix template req_id like AUTO-<file>-<chunk>-001 + ensure source ----
            fixed = []
            for k, r in enumerate(reqs, start=1):
                rid = (r.get("req_id") or "").strip()

                # LLM sometimes returns template literally
                if (not rid) or ("<file>" in rid) or ("<chunk>" in rid) or rid.startswith("AUTO-<"):
                    r["req_id"] = f"AUTO-{name}-{i}-{k:03d}"

                # Ensure source exists and is correct
                file_from_chunk = None
                file_match = re.search(r'\[FILE:([^\]]+)\]', chunk)
                if file_match:
                    file_from_chunk = file_match.group(1)

                if not isinstance(r.get("source"), dict):
                    r["source"] = {}
                current_file = r["source"].get("file") or ""
                if not current_file or current_file in ("", "AUTO", "unknown", "UnknownFile", "<file>", "{file_name}"):
                    r["source"]["file"] = file_from_chunk or name
                if r["source"].get("chunk") is None:
                    r["source"]["chunk"] = i

                # --- xlsx source: extract file, sheet name and row number ---
                if suffix in (".xlsx", ".xls"):
                    parsed_rows = _parse_sheet_rows(chunk)
                    if parsed_rows:
                        req_text = (r.get("requirement") or "").strip()
                        matched_file, sheet, row = _match_sheet_row(req_text, parsed_rows)
                        if matched_file:
                            r["source"]["file"] = matched_file
                        if sheet:
                            r["source"]["sheet"] = sheet
                        if row >= 0:
                            r["source"]["row"] = row

                # --- docx IBM Matrix: force req_id from ROW_ID ---
                if suffix == ".docx" and "[ROW_ID=" in chunk:
                    cur_rid = (r.get("req_id") or "").strip()
                    if not cur_rid.startswith("RFQ-"):
                        row_id_matches = re.findall(r'\[ROW_ID=([^\]]+)\]', chunk)
                        if row_id_matches:
                            for raw_id in row_id_matches:
                                raw_id = raw_id.strip()
                                m = re.match(r'^(HOST|BIOS|BMC|MECH|ENV|SER|AIC|MB|SYS)[\s\-](\d+)$', raw_id, re.IGNORECASE)
                                if m:
                                    prefix = m.group(1).upper()
                                    num = int(m.group(2))
                                    r["req_id"] = f"RFQ-{prefix}-{num:03d}"
                                    break
                    if not (r.get("req_id") or "").startswith("RFQ-"):
                        item_match = re.search(r'Item:\s*(HOST|BIOS|BMC)\s+(\d+)', chunk, re.IGNORECASE)
                        if item_match:
                            prefix = item_match.group(1).upper()
                            num = int(item_match.group(2))
                            r["req_id"] = f"RFQ-{prefix}-{num:03d}"

                # --- 從 chunk 解析位置資訊（table/row/sheet/pdf page）---
                if not isinstance(r.get("source"), dict):
                    r["source"] = {}

                # docx table: [TABLE N ROW N][ROW_ID=XXX] — per-requirement 匹配
                row_id_map = {}
                for _m in re.finditer(r'\[TABLE\s+(\d+)\s+ROW\s+(\d+)\][^\[]*\[ROW_ID=([^\]]+)\]', chunk):
                    _ti, _ri, _rid = int(_m.group(1)), int(_m.group(2)), _m.group(3).strip()
                    row_id_map[_rid] = (_ti, _ri)
                    row_id_map[_rid.replace(' ', '-')] = (_ti, _ri)

                req_id_str = str(r.get("req_id") or "")
                matched_pos = None

                m_rfq = re.match(r'^RFQ-([A-Z]+)-(\d+)$', req_id_str)
                if m_rfq:
                    prefix = m_rfq.group(1)
                    num = int(m_rfq.group(2))
                    for candidate in [f"{prefix} {num}", f"{prefix}-{num}", f"{prefix}{num}"]:
                        if candidate in row_id_map:
                            matched_pos = row_id_map[candidate]
                            break

                if matched_pos:
                    r["source"]["table"] = matched_pos[0]
                    r["source"]["table_row"] = matched_pos[1]
                elif row_id_map:
                    first = next(iter(row_id_map.values()))
                    r["source"]["table"] = first[0]
                    r["source"]["table_row"] = first[1]
                else:
                    tbl_match = re.search(r'\[TABLE\s+(\d+)\s+ROW\s+(\d+)\]', chunk)
                    if tbl_match:
                        r["source"]["table"] = int(tbl_match.group(1))
                        r["source"]["table_row"] = int(tbl_match.group(2))

                # xlsx: [FILE:xxx SHEET:yyy ROW N] — 已有 sheet/row，確認保留
                sheet_match = re.search(r'\[FILE:[^\]]*\s+SHEET:([^\s\]]+)\s+ROW\s+(\d+)\]', chunk)
                if sheet_match and not r["source"].get("sheet"):
                    r["source"]["sheet"] = sheet_match.group(1)
                    r["source"]["row"] = int(sheet_match.group(2))

                # pdf page: [PDF PAGE N]
                pdf_match = re.search(r'\[PDF PAGE\s+(\d+)\]', chunk)
                if pdf_match and not r["source"].get("page"):
                    r["source"]["page"] = int(pdf_match.group(1))

                # --- add traceability fields ---
                r["source_short"] = f"{name}#chunk{i}"
                r["excerpt"] = (chunk[:400] + " ...") if len(chunk) > 400 else chunk
                # keep notes for debugging excerpt origin
                r["notes"] = (r.get("notes", "") + " | HAS_EXCERPT").strip(" |")

                fixed.append(r)

            reqs = fixed

            # Append to in-memory list
            all_reqs.extend(reqs)

            # Save partial per chunk (resume support)
            append_partial(partial_path, name, i, reqs)

    # Deduplicate in case of resume reruns
    all_reqs = dedup_requirements(all_reqs)

    req_doc = {
        "meta": {
            "doc_name": "llm_extracted",
            "case_id": case_id,
            "extracted_at": now_iso(),
            "model": args.model,
            "file_count": len(files),
        },
        "requirements": all_reqs,
    }

    out_path = out_dir / "requirements.json"
    out_path.write_text(json.dumps(req_doc, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] Output: {out_path}")
    if failed_chunks:
        print(f"[OK] Requirements count: {len(all_reqs)} "
              f"(skipped {len(failed_chunks)} chunk(s))")
    else:
        print(f"[OK] Requirements count: {len(all_reqs)}")
    print(f"[OK] Partial saved: {partial_path}")
    if failed_chunks:
        print(f"[WARN] {len(failed_chunks)} chunk(s) failed extraction; "
              f"see {errors_path}")


if __name__ == "__main__":
    main()